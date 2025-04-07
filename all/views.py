from rest_framework import viewsets, permissions, status
from rest_framework.decorators import action
from rest_framework.response import Response
from .models import Object, Apartment, User, ExpenseType, Supplier, Expense, Payment, Document as DocumentModel, UserPayment, SupplierPayment
from .serializers import (ObjectSerializer, ApartmentSerializer, UserSerializer,
                         ExpenseTypeSerializer, SupplierSerializer, ExpenseSerializer,
                         PaymentSerializer, UserPaymentSerializer, DocumentSerializer, SupplierPaymentSerializer)
from .pagination import CustomPagination
from rest_framework_simplejwt.views import TokenObtainPairView
from rest_framework_simplejwt.serializers import TokenObtainPairSerializer
from django.http import FileResponse
from docx import Document
import os
from django.conf import settings
from django.db.models import Sum, Count, Avg
from datetime import datetime
from decimal import Decimal

class ObjectViewSet(viewsets.ModelViewSet):
    queryset = Object.objects.all()
    serializer_class = ObjectSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['name', 'floors', 'total_apartments']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class ApartmentViewSet(viewsets.ModelViewSet):
    queryset = Apartment.objects.all()
    serializer_class = ApartmentSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = [
        'object',
        'rooms',
        'floor',
        'status',
        'price',
        'area',  # Maydon bo'yicha filtr
        'room_number'  # Xonadon raqami bo'yicha qidiruv
    ]

    def get_queryset(self):
        queryset = super().get_queryset()
        # Qo'shimcha filtrlar
        min_price = self.request.query_params.get('min_price', None)
        max_price = self.request.query_params.get('max_price', None)
        min_area = self.request.query_params.get('min_area', None)
        max_area = self.request.query_params.get('max_area', None)

        # Narx bo'yicha filtr
        if min_price:
            queryset = queryset.filter(price__gte=min_price)
        if max_price:
            queryset = queryset.filter(price__lte=max_price)

        # Maydon bo'yicha filtr
        if min_area:
            queryset = queryset.filter(area__gte=min_area)
        if max_area:
            queryset = queryset.filter(area__lte=max_area)

        return queryset

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['user_type', 'phone_number', 'balance']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAdminUser])
    def add_balance(self, request, pk=None):
        user = self.get_object()
        amount = request.data.get('amount', 0)
        try:
            user.add_balance(float(amount))
            return Response({'message': f"{user.fio} balansiga {amount} so‘m qo‘shildi", 'balance': user.balance})
        except ValueError as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class ExpenseTypeViewSet(viewsets.ModelViewSet):
    queryset = ExpenseType.objects.all()
    serializer_class = ExpenseTypeSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['name']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class SupplierViewSet(viewsets.ModelViewSet):
    queryset = Supplier.objects.all()
    serializer_class = SupplierSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['company_name', 'phone_number']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class ExpenseViewSet(viewsets.ModelViewSet):
    queryset = Expense.objects.all()
    serializer_class = ExpenseSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['date', 'supplier', 'expense_type', 'object', 'status']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class PaymentViewSet(viewsets.ModelViewSet):
    queryset = Payment.objects.all()
    serializer_class = PaymentSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['user', 'apartment', 'payment_type', 'created_at', 'status']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def download_contract(self, request, pk=None):
        payment = self.get_object()
        user = payment.user
        apartment = payment.apartment
        obj = apartment.object

        doc = Document()
        doc.add_heading(f"ДАСТЛАБКИ ШАРТНОМА № {payment.id}", 0)
        doc.add_paragraph("Куп хонадонли турар-жой биноси куриш ва сотиш тугрисида")
        doc.add_paragraph(f"« 05 » Февраль 2025 йил\tҚўқон шаҳри")
        doc.add_paragraph(
            f"Қўқон шаҳар «AXLAN HOUSE» МЧЖ номидан низомга асосан фаолият юритувчи раҳбари SODIQOV XASANJON MUXSINJONOVICH "
            f"(кейинги уринларда-«Бажарувчи» деб юритилади) бир томондан ҳамда {user.fio} (келгусида «Куп хонадонли турар-жой биносининг хонадон эгаси-Буюртмачи» деб аталади) "
            f"иккинчи томондан Ўзбекистон Республикасининг «Хужалик юритувчи субъектлар фаолиятининг шартномавий-хуқуқий базаси туғрисида»ги қонунига мувофиқ мазкур шартномани қуйидагилар туғрисида туздик."
        )

        doc.add_heading("ШАРТНОМА ПРЕДМЕТИ.", level=1)
        doc.add_paragraph(
            f"1. Томонлар «Буюртмачи» хонадон сотиб олишга розилиги туғрисида «Бажарувчи» га ариза орқали мурожаат этилгандан сўнг, "
            f"Ўзбекистон Республикаси, Ф appellants вилояти, Қўқон шаҳар {obj.address} да жойлашган {obj.floors} қаватли {obj.total_apartments} хонадонли "
            f"{apartment.room_number}-хонадонli турар-жой биносини қуришга, буюртмачи вазифасини бажариш тўғрисида шартномани (кейинги уринларда - асосий шартнома) тузиш мажбуриятини ўз зиммаларига оладилар."
        )

        doc.add_heading("МУҲИМ ШАРТЛАР.", level=1)
        doc.add_paragraph(
            f"а) «Буюртмачи»га топшириладиган уйнинг {apartment.room_number}-хонадон ({apartment.rooms}-хонали умумий фойдаланиш майдони {apartment.area} кв м) "
            f"умумий қийматининг бошланғич нархи {payment.total_amount} сўмни ташкил этади ва ушбу нарх томонлар томонидан келишилган ҳолда ўзгариши мумкин;"
        )
        doc.add_paragraph(
            f"б) Бажарувчи «тайёр ҳолда топшириш» шартларида турар-жой биносини қуришга бажарувчи вазифасини бажариш мажбуриятини ўз зиммасига олади..."
        )

        doc.add_heading("ХИСОБ-КИТОБ ТАРТИБИ.", level=1)
        doc.add_paragraph(
            f"«Буюртмачи» томонидан мазкур шартнома имзолангач {payment.duration_months} ой давомида яъни 31.12.2025 йилга қадар "
            f"хонадон қуришга пул ўтказиш йўли орқали банкдаги ҳисоб-варағига хонадон қийматининг 100 фоизи яъни {payment.total_amount} сўм миқдорида пул маблағини ўтказади."
        )
        if payment.payment_type == 'muddatli':
            doc.add_paragraph(
                f"Бошланғич тўлов: {payment.initial_payment} сўм, Фоиз: {payment.interest_rate}%, Ҳар ойлик тўлов: {payment.monthly_payment} сўм."
            )
        elif payment.payment_type == 'band':
            doc.add_paragraph(
                f"Band qilish uchun to‘lov: {payment.initial_payment} so‘m, Muddat: {payment.reservation_deadline}."
            )

        docx_path = os.path.join(settings.MEDIA_ROOT, f"contracts/docx/contract_{payment.id}.docx")
        os.makedirs(os.path.dirname(docx_path), exist_ok=True)
        doc.save(docx_path)

        DocumentModel.objects.create(
            payment=payment,
            docx_file=f"contracts/docx/contract_{payment.id}.docx",
            pdf_file=None
        )

        return FileResponse(open(docx_path, 'rb'), as_attachment=True, filename=f"contract_{payment.id}.docx")

    @action(detail=False, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def statistics(self, request):
        today = datetime.now().date()
        total_sales = Payment.objects.aggregate(total=Sum('total_amount'))['total'] or Decimal('0')
        sold_apartments = Apartment.objects.filter(status='sotilgan').count()
        clients = User.objects.filter(user_type='mijoz').count()
        total_objects = Object.objects.count()
        total_apartments = Apartment.objects.count()
        free_apartments = Apartment.objects.filter(status='bosh').count()
        reserved_apartments = Apartment.objects.filter(status='band').count()
        average_price = Apartment.objects.aggregate(avg=Avg('price'))['avg'] or Decimal('0')
        total_payments = Payment.objects.aggregate(total=Sum('monthly_payment'))['total'] or Decimal('0')
        paid_payments = Payment.objects.filter(status='paid').aggregate(total=Sum('paid_amount'))['total'] or Decimal('0')
        pending_payments = Payment.objects.filter(status='pending').aggregate(total=Sum('monthly_payment'))['total'] or Decimal('0')
        overdue_payments = Payment.objects.filter(status='overdue').aggregate(total=Sum('monthly_payment'))['total'] or Decimal('0')
        payments_due_today = Payment.objects.filter(due_date=today.day, status='pending').aggregate(total=Sum('monthly_payment'))['total'] or Decimal('0')
        payments_paid_today = Payment.objects.filter(status='paid', created_at__date=today).aggregate(total=Sum('paid_amount'))['total'] or Decimal('0')

        data = {
            'total_sales': total_sales,
            'sold_apartments': sold_apartments,
            'clients': clients,
            'total_objects': total_objects,
            'total_apartments': total_apartments,
            'free_apartments': free_apartments,
            'reserved_apartments': reserved_apartments,
            'average_price': average_price,
            'total_payments': total_payments,
            'paid_payments': paid_payments,
            'pending_payments': pending_payments,
            'overdue_payments': overdue_payments,
            'payments_due_today': payments_due_today,
            'payments_paid_today': payments_paid_today,
        }
        return Response(data)

class UserPaymentViewSet(viewsets.ModelViewSet):
    queryset = UserPayment.objects.all()
    serializer_class = UserPaymentSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['user', 'date', 'payment_type']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class SupplierPaymentViewSet(viewsets.ModelViewSet):
    queryset = SupplierPayment.objects.all()
    serializer_class = SupplierPaymentSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['supplier', 'date', 'payment_type']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class DocumentViewSet(viewsets.ModelViewSet):
    queryset = DocumentModel.objects.all()
    serializer_class = DocumentSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['payment', 'created_at', 'document_type']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class CustomTokenObtainPairSerializer(TokenObtainPairSerializer):
    @classmethod
    def get_token(cls, user):
        token = super().get_token(user)
        token['user_type'] = user.user_type
        token['fio'] = user.fio
        return token

class CustomTokenObtainPairView(TokenObtainPairView):
    serializer_class = CustomTokenObtainPairSerializer